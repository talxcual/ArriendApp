import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def main():
    desktop_path = r"C:\Users\toled\Desktop\Documentacion_LuxeRental"
    if not os.path.exists(desktop_path):
        os.makedirs(desktop_path)
        
    doc_path = os.path.join(desktop_path, "Documentacion_LuxeRental_Pro.docx")
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Documentación Oficial: LuxeRental Pro', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Generado por el Analista de Sistemas y Technical Writer', style='Subtitle')
    
    # 1. Informe Profesional del Proyecto
    doc.add_heading('1. Informe Profesional del Proyecto (Resumen Ejecutivo)', level=1)
    
    doc.add_heading('Descripción General', level=2)
    doc.add_paragraph('LuxeRental Pro es una solución de software móvil diseñada para optimizar y centralizar la gestión operativa de empresas dedicadas al arriendo de equipos y mobiliario. El sistema resuelve el problema común de la pérdida de información, la descoordinación de fechas y la ineficiencia en la comunicación con los clientes. Mediante una plataforma unificada, el personal puede administrar su inventario, agendar despachos precisos y mantener un seguimiento claro de cada servicio brindado, mejorando significativamente la productividad y la satisfacción del usuario final.')
    
    doc.add_heading('Alcance del Proyecto', level=2)
    p = doc.add_paragraph('La aplicación contempla los siguientes módulos y funcionalidades principales:\n')
    p.add_run('- Gestión de Inventario: ').bold = True
    p.add_run('Visualización de catálogo en tiempo real con estados (Disponible, Ocupado) y precios dinámicos.\n')
    p.add_run('- Programación de Reservas (Calendario): ').bold = True
    p.add_run('Asignación de equipos a clientes específicos, con selectores de fecha y hora para el evento.\n')
    p.add_run('- Registro y Seguimiento de Clientes: ').bold = True
    p.add_run('Almacenamiento de datos clave (Nombre, Dirección, Teléfono) vinculados a la orden.\n')
    p.add_run('- Comunicaciones Integradas: ').bold = True
    p.add_run('Acceso de un solo clic para contactar al cliente vía llamada telefónica tradicional o mediante WhatsApp, directamente desde el detalle de la reserva.')
    
    doc.add_heading('Stack Tecnológico', level=2)
    p2 = doc.add_paragraph()
    p2.add_run('El sistema está desarrollado íntegramente de manera nativa para dispositivos Android, garantizando un rendimiento óptimo y una integración profunda con los servicios del sistema operativo:\n')
    p2.add_run('- Lenguaje Base: ').bold = True
    p2.add_run('Kotlin, bajo el paradigma de programación reactiva y corrutinas.\n')
    p2.add_run('- Interfaz de Usuario (UI): ').bold = True
    p2.add_run('Jetpack Compose para componentes dinámicos y fluidos, respaldado por XML clásico para configuraciones estructurales.\n')
    p2.add_run('- Backend/BaaS: ').bold = True
    p2.add_run('Firebase (Firestore para la base de datos en tiempo real y Authentication para seguridad de acceso).\n')
    p2.add_run('- Arquitectura: ').bold = True
    p2.add_run('Patrón MVVM (Model-View-ViewModel).')

    # 2. Documentación Técnica
    doc.add_page_break()
    doc.add_heading('2. Documentación Técnica (Arquitectura y Sistemas)', level=1)
    
    doc.add_heading('Modelo de Datos', level=2)
    doc.add_paragraph('El ecosistema se apoya en tres entidades lógicas principales, diseñadas para operar eficientemente a través del repositorio:')
    p3 = doc.add_paragraph()
    p3.add_run('- UserEntity (Operador): ').bold = True
    p3.add_run('Administra los accesos de los trabajadores de la empresa. Atributos: UID, nombre, rol.\n')
    p3.add_run('- InventoryEntity (Equipo): ').bold = True
    p3.add_run('Representa los ítems físicos disponibles para arriendo. Atributos: ID, nombre, imagen URL, precio (CLP), estado actual.\n')
    p3.add_run('- ReservaEntity (Reserva/Transacción): ').bold = True
    p3.add_run('Unifica el inventario con el cliente. Atributos: ID, equipmentId, fechaHora (momento del evento), fechaRegistro (timestamp de creación del registro), clienteNombre, clienteDireccion, clienteContacto, monto total, estado operativo y workerId (trazabilidad de quién tomó el pedido).')
    
    doc.add_heading('Lógica de Negocio y Flujos', level=2)
    doc.add_paragraph('El proceso central de negocio sigue un flujo transaccional controlado:\n'
                      '1. Selección: El operador explora el catálogo de equipos filtrando aquellos en estado "DISPONIBLE".\n'
                      '2. Captura de Datos: A través del ViewModel, se inicializa el formulario solicitando al operador los datos demográficos del cliente y las fechas mediante los Date/Time Pickers nativos.\n'
                      '3. Commit (Confirmación): Se genera una transacción que (a) inserta la nueva Reserva en la colección de la base de datos y (b) muta el estado del Equipo asociado a "OCUPADO".\n'
                      '4. Estados Dinámicos: La UI (Compose) observa el StateFlow de las entidades. Cualquier modificación se renderiza inmediatamente, moviendo el ítem reservado a las vistas de seguimiento y calendario.')
    
    doc.add_heading('Integraciones (Intents e Intercomunicación)', level=2)
    doc.add_paragraph('LuxeRental Pro hace uso de Intents (explícitos e implícitos) de Android para romper el aislamiento del sistema y extender funcionalidades sin duplicar herramientas:')
    p4 = doc.add_paragraph()
    p4.add_run('- Llamadas Telefónicas (ACTION_DIAL): ').bold = True
    p4.add_run('Se inyecta un Intent implícito con la URI `tel:$numeroContacto`. El sistema operativo lo delega al marcador (dialer) preferido del usuario.\n')
    p4.add_run('- Integración con WhatsApp (ACTION_VIEW): ').bold = True
    p4.add_run('Se implementa un Intent explícito forzando el paquete destino (`intent.setPackage("com.whatsapp")`). Utilizando la API web de WhatsApp (`api.whatsapp.com/send?phone=$numero`), la app puentea directamente al chat del cliente sin necesidad de agregarlo previamente a la libreta de contactos. Se gestiona la excepción `ActivityNotFoundException` como medida preventiva de seguridad.')

    # 3. Manual de Usuario
    doc.add_page_break()
    doc.add_heading('3. Manual de Usuario (Guía Operativa)', level=1)
    
    doc.add_paragraph('¡Bienvenido a LuxeRental Pro! Este manual te guiará paso a paso para que puedas aprovechar al máximo tu nueva herramienta de gestión diaria.')
    
    doc.add_heading('A. ¿Cómo registrar un nuevo arriendo?', level=2)
    doc.add_paragraph('Sigue estos pasos cuando un cliente te confirme la contratación de un servicio:')
    doc.add_paragraph('1. Entra a la pestaña Catálogo desde la pantalla principal.')
    doc.add_paragraph('2. Busca el equipo solicitado. Asegúrate de que tenga una etiqueta verde que diga "DISPONIBLE".')
    doc.add_paragraph('3. Toca la tarjeta del equipo. Se desplegará el panel de "Confirmar Nuevo Arriendo".')
    
    p_img1 = doc.add_paragraph('[Insertar Captura de Pantalla Aquí: Panel de Formulario de Nuevo Arriendo]')
    p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('4. Ingresa el Nombre del Cliente, la Dirección de Entrega y el Teléfono de Contacto (recomendable usar el código de país, ej. 569...).')
    doc.add_paragraph('5. Haz clic en "Fecha" y selecciona el día del evento en el calendario que aparecerá. Luego, presiona "Hora" y establece el horario de entrega.')
    doc.add_paragraph('6. Presiona Confirmar Arriendo. ¡Listo! El equipo se marcará como ocupado y el arriendo quedará registrado en tu agenda.')

    doc.add_heading('B. ¿Cómo visualizar el calendario de reservas?', level=2)
    doc.add_paragraph('Tu agenda operativa en tiempo real:')
    doc.add_paragraph('1. Dirígete a la pantalla de Calendario.')
    doc.add_paragraph('2. Aquí verás el listado cronológico de todas las reservas, mostrando un vistazo rápido al estado (ej. Pendiente, Entregado) y la fecha de los eventos programados.')
    
    p_img2 = doc.add_paragraph('[Insertar Captura de Pantalla Aquí: Pantalla de Calendario con Lista de Reservas]')
    p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('C. ¿Cómo contactar al cliente desde una reserva?', level=2)
    doc.add_paragraph('Para enviar ubicaciones, cobrar saldos o verificar instrucciones de entrega rápidamente:')
    doc.add_paragraph('1. Estando en el Calendario, pulsa sobre la tarjeta de la reserva que deseas gestionar.')
    doc.add_paragraph('2. Se abrirá una tarjeta grande con todos los "Detalles de Reserva". En la parte inferior, verás dos botones grandes: "Llamar" y "WhatsApp".')
    
    p_img3 = doc.add_paragraph('[Insertar Captura de Pantalla Aquí: Modal de Detalles de Reserva con botones de contacto]')
    p_img3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('3. Botón Llamar: Automáticamente abrirá el teléfono de tu celular con el número del cliente listo para marcar.\n'
                      '4. Botón WhatsApp: Te llevará de manera directa a un chat nuevo de WhatsApp con el cliente para escribirle de forma inmediata, ¡sin tener que añadirlo como contacto en tu teléfono!')
    
    final_p = doc.add_paragraph()
    final_p.add_run('¡Con estos pasos básicos ya estás listo para llevar el control absoluto de tus operaciones!').italic = True

    doc.save(doc_path)

if __name__ == "__main__":
    main()
